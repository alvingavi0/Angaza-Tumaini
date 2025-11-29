from docx import Document
from docx.shared import Pt

doc = Document()

# Title
h = doc.add_heading('Angaza Tumaini Mission Center — Project Documentation', level=1)

# Project Overview
doc.add_heading('1. Project Overview', level=2)
doc.add_paragraph("Angaza Tumaini is a static website built with plain HTML, Tailwind CSS (via CDN), and small client-side JavaScript. The site includes pages: index.html (home), contact.html (contact), newsletter.html (newsletter & socials). Local image assets and PNG icons live in the folder filez/ and filez/icons/.")

# Timeline
doc.add_heading('2. Development timeline (Beginning → End)', level=2)
steps = [
    'Initial scaffolding: Static HTML pages for Home, Contact, and supporting assets. Tailwind loaded from CDN and Lucide icons used for vector icons initially.',
    'Iterative updates: Footer social icons were added and changed from inline/SVG to local PNG files to ensure consistent rendering on user systems.',
    'Newsletter page: Created newsletter.html and wired an in-site CTA that navigates from footer button to this page.',
    'Contact updates: Mailto and tel links set to AngazaTumaini.org@gmail.com and tel:+254716475764, WhatsApp short link https://wa.link/xjt1s4 added.',
    'Mailto fallback: Small JS heuristic added to open Gmail compose if mailto doesn\'t open a native client.',
    'Visual/UX fixes: Circular PNG icons for social links, header/nav restored to original layout after accidental change.',
    'Deployment: Deployed to Vercel using the CLI. Repository was committed and pushed to GitHub. Cleaned up unused SVG icon files from filez/icons/ and pushed the cleanup commit.'
]
for s in steps:
    doc.add_paragraph('- ' + s)

# Files changed
doc.add_heading('3. Files changed / created', level=2)
doc.add_paragraph('- index.html: Footer icons, mailto/tel links, mailto fallback script, header restored, WhatsApp icon added.')
doc.add_paragraph('- contact.html: Updated contact email, phone, and WhatsApp link; contact form using Formspree with mailto fallback.')
doc.add_paragraph('- newsletter.html: New page created listing social channels and a mailto subscribe link.')
doc.add_paragraph('- filez/icons/: PNG icons added (facebook.png, Instagram.png, Gmail.png, tiktok.png, Whatsapp.png). Unused SVGs removed.')

# Local preview
doc.add_heading('4. How to preview locally', level=2)
doc.add_paragraph('Open a terminal in the project root and run a simple static server. Example using Python 3:')
doc.add_paragraph('python -m http.server 8000')
doc.add_paragraph('Then open http://127.0.0.1:8000/index.html in your browser to view the site.')

# Deployment
doc.add_heading('5. Deployment', level=2)
doc.add_paragraph('- Vercel: Deployed via `npx vercel --prod` from project root. If CLI rate limits occur, connect the GitHub repo in the Vercel dashboard to enable Git-based auto-deploy on push.')

# Screenshots placeholders
doc.add_heading('6. Screenshots (placeholders)', level=2)
doc.add_paragraph('Page: index.html')
doc.add_paragraph('[Insert screenshot of index.html here — paste image below this line]')
doc.add_paragraph('')
doc.add_paragraph('Page: contact.html')
doc.add_paragraph('[Insert screenshot of contact.html here — paste image below this line]')
doc.add_paragraph('')
doc.add_paragraph('Page: newsletter.html')
doc.add_paragraph('[Insert screenshot of newsletter.html here — paste image below this line]')

# Billing
doc.add_heading('7. Billing', level=2)
doc.add_paragraph('Service: Web development and deployment for Angaza Tumaini Mission Center')
doc.add_paragraph('Amount: 21,000 shillings')
doc.add_paragraph('Payment via phone number: 0759106034')

# Notes
doc.add_heading('8. Notes & Next steps', level=2)
doc.add_paragraph('- Verify live site on Vercel and test mailto/tel/WhatsApp links on target devices. Provide screenshots and any requested copy changes and I will update and redeploy.')
doc.add_paragraph('- Optional: wire the newsletter subscription to a mailing provider (Mailchimp, ConvertKit, MailerLite) or add a server-side endpoint to collect emails.')

doc.add_page_break()

doc.save('../Angaza-Tumaini-Documentation.docx')
print('Saved Angaza-Tumaini-Documentation.docx')
